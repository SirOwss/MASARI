import gc
import time
import camelot
import matplotlib.pyplot as plt
import io
import re
import pandas as pd
from datetime import date, datetime, timedelta, time
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
import numpy as np
import random

# --- Utility Functions ---

def arabic_to_western(s):
    if pd.isna(s):
        return s
    arabic_digits = '٠١٢٣٤٥٦٧٨٩'
    western_digits = '0123456789'
    trans_table = str.maketrans(arabic_digits, western_digits)
    return s.translate(trans_table)

def clean_count(df):
    # Ensure 'registeredCount' exists and is string type for processing
    if 'registeredCount' not in df.columns:
        df['registeredCount'] = np.nan # Or handle error appropriately
        return df

    df['registeredCount'] = df['registeredCount'].astype(str)

    # --- Pre-processing: Standardize 'empty' cells to NaN ---
    # Replace empty strings or strings with only spaces with NaN
    df['registeredCount'] = df['registeredCount'].replace(r'^\s*$', np.nan, regex=True)

    # Create a temporary column to store the cleaned numbers
    df['cleaned_number'] = np.nan

    # --- First Pass: Handle 3-part splits and directly assigned single numbers ---
    for i, row in df.iterrows():
        val = str(row['registeredCount']).strip()

        if pd.isna(val) or val == 'nan':
            continue

        parts = val.split()

        if len(parts) == 1:
            df.loc[i, 'cleaned_number'] = parts[0]
        elif len(parts) == 3:
            df.loc[i, 'cleaned_number'] = parts[1]

            if i > 0 and (pd.isna(df.loc[i-1, 'cleaned_number']) or str(df.loc[i-1, 'cleaned_number']).strip() == ''):
                df.loc[i-1, 'cleaned_number'] = parts[0]

            if i < len(df) - 1 and (pd.isna(df.loc[i+1, 'cleaned_number']) or str(df.loc[i+1, 'cleaned_number']).strip() == ''):
                df.loc[i+1, 'cleaned_number'] = parts[2]

    # --- Second Pass: Handle 2-part splits and remaining empty cells ---
    df['temp_number'] = df['cleaned_number'].copy()

    max_iterations = 5
    for _ in range(max_iterations):
        changes_made = False
        for i, row in df.iterrows():
            original_val_in_cell = str(df.loc[i, 'registeredCount']).strip()

            if original_val_in_cell and len(original_val_in_cell.split()) == 2:
                current_cleaned_val = df.loc[i, 'temp_number']

                if pd.isna(current_cleaned_val) or str(current_cleaned_val).strip() == '':
                    parts = original_val_in_cell.split()
                    assigned_current = False

                    if i > 0 and (pd.isna(df.loc[i-1, 'temp_number']) or str(df.loc[i-1, 'temp_number']).strip() == ''):
                        df.loc[i-1, 'temp_number'] = parts[0]
                        df.loc[i, 'temp_number'] = parts[1]
                        assigned_current = True
                        changes_made = True
                    elif i < len(df) - 1 and (pd.isna(df.loc[i+1, 'temp_number']) or str(df.loc[i+1, 'temp_number']).strip() == ''):
                        df.loc[i, 'temp_number'] = parts[0]
                        df.loc[i+1, 'temp_number'] = parts[1]
                        assigned_current = True
                        changes_made = True

                    if not assigned_current:
                        df.loc[i, 'temp_number'] = parts[0]

        df['cleaned_number'] = df['temp_number'].copy()
        if not changes_made:
            break

    # --- Final Cleanup ---
    df['cleaned_number'] = pd.to_numeric(df['cleaned_number'], errors='coerce')

    df.drop(columns=['registeredCount', 'temp_number'], inplace=True)
    df.rename(columns={'cleaned_number': 'registeredCount'}, inplace=True)
    return df

def get_earliest_day(days):
    if pd.isna(days):
        return None
    for d in days_order:
        if d in days:
            return d
    return None

def extract_start_time1(timeslot):
    timeslot = arabic_to_western(timeslot)
    if '-' in timeslot:
        return timeslot.split('-')[1].strip()
    return None

def extract_start_time(timeslot):
    if pd.isna(timeslot):
        return time(23, 59)

    timeslot = arabic_to_western(str(timeslot))

    if '-' in timeslot:
        time_str = timeslot.split('-')[1].strip()
    else:
        time_str = timeslot.strip()

    if re.fullmatch(r'\d{4}', time_str):
        time_str = time_str[:2] + ':' + time_str[2:]

    try:
        return datetime.strptime(time_str, '%H:%M').time()
    except ValueError:
        return time(23, 59)

def normalize_course_name(name):
    match = re.match(r"(\d+)-([A-Z]+)", name)
    if match:
        number, dept = match.groups()
        return f"{dept}-{number}"
    return name

def get_day_priority(days):
    if pd.isna(days):
        return float('inf')
    for d in day_order:
        if d in days:
            return day_order[d]
    return float('inf')

def set_cell_shading(cell, color_hex):
    """
    Sets the background shading color for a given table cell.
    Args:
        cell: The docx.table._Cell object.
        color_hex (str): The hexadecimal color code (e.g., "D3D3D3" for light gray).
    """
    tcPr = cell._element.get_or_add_tcPr()
    tcShd = OxmlElement("w:shd")
    tcShd.set(qn("w:fill"), color_hex)
    tcPr.append(tcShd)

def parse_gregorian_date(date_str_list):
    """
    Parses a Gregorian date from a list, handling multiple formats.
    Prioritizes the last valid date string in the list.
    Returns a datetime object.
    """
    if not isinstance(date_str_list, list) or not date_str_list:
        return None

    for date_str in reversed(date_str_list):
        if date_str is None:
            continue
        try:
            # Try 2024-MM-DD format
            return datetime.strptime(date_str, "%Y-%m-%d")
        except ValueError:
            try:
                # Try DD/MM/YYYY format
                return datetime.strptime(date_str, "%d/%m/%Y")
            except ValueError:
                # If both formats fail, log the error and continue to the next string
                # print(f"Failed to parse date string '{date_str}' with known formats.") # Uncomment for detailed parsing errors
                continue
    return None # Return None if no date string could be parsed

def extract_islamic_date_from_list(date_str_list, gregorian_date_obj):
    """
    Extracts the Islamic date string from the provided list.
    It attempts to find a date string that is not the Gregorian date.
    Assumes Islamic date is usually the first non-Gregorian date-like string.
    """
    if not isinstance(date_str_list, list) or not date_str_list:
        return None

    gregorian_str_ymd = gregorian_date_obj.strftime("%Y-%m-%d") if gregorian_date_obj else None
    gregorian_str_dmy = gregorian_date_obj.strftime("%d/%m/%Y") if gregorian_date_obj else None

    for date_str in date_str_list:
        if date_str is None:
            continue
        # Check if the string matches the Gregorian date in either format
        if date_str == gregorian_str_ymd or date_str == gregorian_str_dmy:
            continue # Skip Gregorian date strings

        # If it's a string and not the Gregorian date, assume it's the Islamic date
        # You might add more robust checks here if Islamic date formats vary widely
        return date_str
    return None # No distinct Islamic date found

def create_exam_schedule_docx(dataframe): # Removed filename, will return BytesIO
    """
    Creates a Word document with an exam schedule table and returns it as a BytesIO object.
    Includes dynamic venue assignment with capacity limits and conflict resolution.
    """
    document = Document()

    document.add_heading("Final Exam Schedule for the Second Semester 1446 H", level=1)

    desired_headers = [
        "S#",
        "Course Name",
        "Instructor(s) (#Students)",
        "Exam Day & Date",
        "Exam Time",
        "Exam Room"
    ]

    time_slot_mapping = {
        1: "09:00 AM to 11:00 AM",
        2: "11:30 AM to 01:30 PM",
        3: "02:00 PM to 04:00 PM"
    }

    # Define available venues and their capacity
    all_venues = [101, 102, 103, 104, 105, 106, 107, 108, 109]
    venue_capacity_per_room = 30

    # This dictionary will track venue occupancy for each date and time slot
    # Key: (date_obj, time_slot_num)
    # Value: {venue_id: current_occupancy_for_this_slot}
    venue_occupancy_tracker = {}

    df_processed = dataframe.copy()

    df_processed['parsed_gregorian_date'] = df_processed['examDate'].apply(parse_gregorian_date)
    df_processed['timeSlot'] = pd.to_numeric(df_processed['timeSlot'], errors='coerce')

    # Handle NaN values in parsed_gregorian_date for sorting and rank calculation
    # Fill NaT with a very large timestamp to ensure they sort to the end
    df_processed['date_group_rank'] = df_processed['parsed_gregorian_date'].fillna(pd.Timestamp.max).rank(method='dense').astype(int)

    df_processed = df_processed.sort_values(
        by=['parsed_gregorian_date', 'timeSlot'],
        na_position='last'
    ).reset_index(drop=True)

    # --- DEBUGGING STEP: Print the sorted DataFrame to verify ---
    #print("--- Sorted DataFrame for Verification ---")
    #print(df_processed[['courseName', 'parsed_gregorian_date', 'timeSlot', 'registeredCount']])
    #print("-----------------------------------------")

    df_processed['islamic_date_str_extracted'] = df_processed.apply(
        lambda row: extract_islamic_date_from_list(row['examDate'], row['parsed_gregorian_date']), axis=1
    )
    df_processed['day_name'] = df_processed['parsed_gregorian_date'].apply(lambda x: x.strftime("%A") if pd.notna(x) else "")

    # --- Venue Assignment Logic ---
    df_processed['assigned_rooms_details'] = None # New column to store assigned room strings

    for idx, row in df_processed.iterrows():
        current_date = row['parsed_gregorian_date']
        current_time_slot = row['timeSlot']
        num_students = row['registeredCount']
        
        # Handle NaN for registeredCount, default to 0
        if pd.isna(num_students):
            num_students = 0
        else:
            num_students = int(num_students)

        teachers = row['teacherName']
        if isinstance(teachers, list):
            teachers_list = [str(name) for name in teachers if name is not None]
        else:
            teachers_list = [str(teachers)] if pd.notna(teachers) else []
        num_teachers = len(teachers_list)

        assigned_rooms_for_exam_lines = []
        students_remaining_for_course = num_students

        # Initialize or get the current occupancy for this date and time slot
        # Ensure the key is hashable (tuple of date and time slot)
        slot_key = (current_date, current_time_slot)
        if slot_key not in venue_occupancy_tracker:
            venue_occupancy_tracker[slot_key] = {venue: 0 for venue in all_venues}
        
        current_slot_occupancy = venue_occupancy_tracker[slot_key]

        # Get rooms that are not fully occupied for this specific date and time slot
        available_venues_for_this_slot = [
            v for v in all_venues
            if current_slot_occupancy[v] < venue_capacity_per_room
        ]
        random.shuffle(available_venues_for_this_slot) # Randomize selection

        rooms_assigned_to_this_course = []
        teacher_index_for_round_robin = 0

        # Try to assign rooms until all students are accommodated or no more rooms are available
        while students_remaining_for_course > 0 and available_venues_for_this_slot:
            selected_venue = available_venues_for_this_slot.pop(0)
            
            # Calculate how many students can fit in this part of the room
            capacity_in_this_part = venue_capacity_per_room - current_slot_occupancy[selected_venue]
            
            # This ensures we don't assign more students than the room can hold or more than needed
            students_to_assign_here = min(students_remaining_for_course, capacity_in_this_part)

            if students_to_assign_here > 0:
                rooms_assigned_to_this_course.append(selected_venue)
                current_slot_occupancy[selected_venue] += students_to_assign_here
                students_remaining_for_course -= students_to_assign_here
            # If students_to_assign_here is 0, this room is full or not needed, skip it
            # No 'else' needed here, loop will continue if students_remaining_for_course > 0

        # Now format the assigned rooms based on teacher count
        if not rooms_assigned_to_this_course:
            assigned_rooms_for_exam_lines.append("No Venue Assigned")
            if students_remaining_for_course > 0:
                assigned_rooms_for_exam_lines.append(f"Not enough capacity for {students_remaining_for_course} students.")
                print(f"Warning: Not enough venue capacity for course {row['courseName']} on {current_date} at {current_time_slot}. {students_remaining_for_course} students unassigned.")
        else:
            if num_teachers == 1:
                # Single teacher: list all assigned rooms with that teacher's name
                room_ids_str = ", ".join(map(str, sorted(rooms_assigned_to_this_course)))
                assigned_rooms_for_exam_lines.append(f"{room_ids_str}: {teachers_list[0]}")
            else:
                # Multiple teachers: assign one teacher per room assigned (round-robin)
                for room_id in sorted(rooms_assigned_to_this_course): # Sort for consistent output
                    assigned_teacher = teachers_list[teacher_index_for_round_robin % num_teachers]
                    assigned_rooms_for_exam_lines.append(f"{room_id}: {assigned_teacher}")
                    teacher_index_for_round_robin += 1
            
            if students_remaining_for_course > 0:
                assigned_rooms_for_exam_lines.append(f"Not enough capacity for {students_remaining_for_course} students.")
                print(f"Warning: Not enough venue capacity for course {row['courseName']} on {current_date} at {current_time_slot}. {students_remaining_for_course} students unassigned.")

        df_processed.at[idx, 'assigned_rooms_details'] = assigned_rooms_for_exam_lines

    # --- Table Creation in DOCX ---
    num_cols = len(desired_headers)
    table = document.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'

    # Set up the header row formatting
    hdr_cells = table.rows[0].cells
    for i, header_name in enumerate(desired_headers):
        hdr_cells[i].text = header_name
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Dark gray text color
        font = hdr_cells[i].paragraphs[0].runs[0].font
        font.size = Pt(10)

    # Set column widths (adjusted for content)
    table.columns[0].width = Inches(0.5)  # S#
    table.columns[1].width = Inches(1.2)  # Course Name
    table.columns[2].width = Inches(2.5)  # Instructor(s) (#Students)
    table.columns[3].width = Inches(1.8)  # Exam Day & Date
    table.columns[4].width = Inches(1.5)  # Exam Time
    table.columns[5].width = Inches(1.5)  # Exam Room (adjusted for multiple rooms and teacher names)

    # Variables to manage cell merging for 'Exam Day & Date'
    table_row_idx_offset = 1 # Accounts for the header row
    last_gregorian_date_for_merge = None

    # Iterate through the processed DataFrame to populate the table rows
    for df_idx, row in df_processed.iterrows():
        new_doc_row = table.add_row()
        row_cells = new_doc_row.cells

        # Apply shading based on date_group_rank
        if row['date_group_rank'] % 2 == 0: # Apply shading to even-numbered date groups
            for cell in row_cells:
                set_cell_shading(cell, "D3D3D3") # Stronger light gray

        current_gregorian_date = row['parsed_gregorian_date']

        # S# (Sequential row number in the DOCX table)
        row_cells[0].text = str(df_idx + 1)

        # Course Name
        row_cells[1].text = str(row["courseName"])

        # Instructor(s) (#Students)
        teacher_names = row["teacherName"]
        if isinstance(teacher_names, list):
            formatted_teachers_display = "\n".join(filter(None, [str(name) for name in teacher_names]))
        else:
            formatted_teachers_display = str(teacher_names)
        registered_count = int(row["registeredCount"]) if pd.notna(row["registeredCount"]) else 0
        row_cells[2].text = f"{formatted_teachers_display}\n({registered_count})"

        # Exam Day & Date - Logic for populating and preparing for merging
        if pd.isna(current_gregorian_date) or current_gregorian_date != last_gregorian_date_for_merge:
            day_name = row['day_name']
            gregorian_date_str = row['parsed_gregorian_date'].strftime("%d/%m/%Y") if pd.notna(row['parsed_gregorian_date']) else ""
            islamic_date_str = row['islamic_date_str_extracted'] if row['islamic_date_str_extracted'] else ""
            row_cells[3].text = f"{day_name}\n{islamic_date_str}\n{gregorian_date_str}"
            last_gregorian_date_for_merge = current_gregorian_date
        else:
            row_cells[3].text = ""

        # Exam Time
        time_slot_num = row["timeSlot"]
        exam_time_str = time_slot_mapping.get(time_slot_num, "N/A")
        row_cells[4].text = exam_time_str

        # Exam Room - Now uses the pre-assigned rooms
        assigned_rooms_list = row['assigned_rooms_details']
        if assigned_rooms_list:
            row_cells[5].text = "\n".join(assigned_rooms_list)
        else:
            row_cells[5].text = "No Venue Assigned"

        # Apply common cell formatting (vertical alignment and font size)
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER # Vertical centering
            if cell.paragraphs and cell.paragraphs[0].runs:
                font = cell.paragraphs[0].runs[0].font
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Dark gray text color
            else:
                run = cell.paragraphs[0].add_run()
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Dark gray text color

            # Ensure horizontal centering for Exam Day & Date column (index 3)
            # This applies to the cell that *contains* the date text (the top-most cell in a merged group)
            if table.cell(table_row_idx_offset, 3) == cell:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


        table_row_idx_offset += 1

    # --- Post-processing: Perform Merging for 'Exam Day & Date' cells ---
    if len(table.rows) > 1:
        current_merge_start_table_row_idx = 1
        current_date_for_merge = df_processed.iloc[0]['parsed_gregorian_date']

        for i in range(1, len(df_processed)):
            table_row_index = i + 1

            next_date_for_merge = df_processed.iloc[i]['parsed_gregorian_date']

            if pd.isna(next_date_for_merge) or next_date_for_merge != current_date_for_merge or (i == len(df_processed) - 1):
                merge_end_row_idx = table_row_index if (pd.notna(next_date_for_merge) and next_date_for_merge == current_date_for_merge) else table_row_index - 1

                if merge_end_row_idx > current_merge_start_table_row_idx:
                    top_cell = table.cell(current_merge_start_table_row_idx, 3)
                    bottom_cell = table.cell(merge_end_row_idx, 3)
                    top_cell.merge(bottom_cell)
                    # Ensure the merged cell content is centered after merging
                    top_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    top_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                if i < len(df_processed) - 1:
                    current_merge_start_table_row_idx = table_row_index
                    current_date_for_merge = next_date_for_merge
                elif i == len(df_processed) - 1 and (pd.isna(next_date_for_merge) or next_date_for_merge != current_date_for_merge):
                    pass

    # Save the document to a BytesIO object instead of a file
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0) # Rewind the buffer to the beginning
    return buffer # Return the BytesIO object



def clean_header(lst):
    filtered = [s for s in lst if re.search(r'[\d]', s)]
    lst = [i.strip('-') for i in filtered]
    if len(lst) == 3:
        return f'{lst[0]}-{lst[1]}-{lst[2]}'
    elif len(lst) == 2:
        d,m = lst[1].split('-')
        return f'{lst[0]}-{m}-{d}'
    return None # Added to handle cases where len is not 2 or 3

def clean_date(s):
    lst = s.split('\n')

    filtered = [s for s in lst if re.search(r'[\d-]', s)]

    parts = []
    for item in filtered:
        item = item.strip('-')
        if '-' in item:
            parts.extend(item.split('-'))
        else:
            parts.append(item)

    def pad_part(x):
        return x if len(x) == 4 else x.zfill(2)

    parts = list(map(pad_part, parts))

    if len(parts) >= 6: # Ensure enough parts for both dates
        hijri_year, hijri_month, hijri_day = parts[:3]
        gregorian_year, gregorian_month, gregorian_day = parts[3:6]

        hijri_date = f"{hijri_year}-{hijri_month}-{hijri_day}"
        gregorian_date = f"{gregorian_year}-{gregorian_month}-{gregorian_day}"
        return hijri_date, gregorian_date
    return None, None # Return None if not enough parts

# --- Global Constants ---
days_order = ['U', 'M', 'T', 'W', 'R']
day_order = {'U': 0, 'M': 1, 'T': 2, 'W': 3, 'R': 4}

# --- Main Script Execution ---
# ... (all imports and utility functions unchanged) ...
def examsScheduling(file1,file2):
    def get_exam_id(row):
        try:
            hour = int(str(row['timeSlot']).split(':')[0])  # get hour from 'HH:MM:SS'
            day = str(row['days']).strip()
        except:
            return None
        
        for eid, slot_hour, slot_day in coursetimings:
            if hour == slot_hour and day == slot_day:
                return eid
        return None
    # PHASE 1: Extract Course IDs and Course Names from tmp33.pdf (pages 1 and 2)
    tables = camelot.read_pdf(file1, pages="1,2", flavor='lattice')

    pg1 = [t for t in tables if t.page == 1]
    pg2 = [t for t in tables if t.page == 2]

    
    courses = []
    valid = ('cocs','coit','cois')

    for i, table in enumerate(pg1):
        for examid, course in zip(table.df[0],table.df[1]):
            if course.lower().startswith(valid):
                courses.append((examid,course))

    # Extract Course Timings from tmp33.pdf (page 2)
    coursetimings = []
    for i, table in enumerate(pg2):
        pervDay = ''
        for examid, examtime, day in zip(table.df[0],table.df[1],table.df[2]):
            IsPm = True
            examtime  = examtime.split('\n')

            valid_times = []
            for item in examtime:
                try:
                    t = datetime.strptime(item.strip(), '%H:%M').time()
                    valid_times.append(t)
                except ValueError:
                    if item[::-1] == 'صباحاً':
                        IsPm = False

            if valid_times:
                smallest_time = min(valid_times)
                dt = datetime.combine(date.today(), smallest_time)
                if IsPm and (smallest_time.hour != 12 and smallest_time.hour < 12): # Corrected AM/PM logic
                    dt += timedelta(hours=12)
                elif not IsPm and smallest_time.hour == 12: # For 12 AM (midnight)
                    dt -= timedelta(hours=12)

                smallest_str = dt.strftime('%H:%M')

            day = day[::-1]

            if 'األحد' in day:
                day = 'U'
                pervDay = day

            elif 'االثنين' in day:
                day = 'M'
                pervDay = day
            elif 'الثالثاء' in day:
                day = 'T'
                pervDay = day
            elif 'األربعاء' in day:
                day = 'W'
                pervDay = day
            elif 'الخميس' in day:
                day = 'R'
                pervDay = day

            coursetimings.append((examid,smallest_str,pervDay))

    coursetimings = [(int(eid), int(hour_str.split(':')[0]), day) for eid, hour_str, day in coursetimings]

    # PHASE 3: Process Main Data Tables from input2.1.pdf

    tables = camelot.read_pdf(file2, pages='all',flavor='lattice')

    dfs = []
    for i, table in enumerate(tables):
        df = table.df

        cols = df.columns.tolist()
        for i,value in enumerate(df.iloc[0]):
            value = value[::-1].strip()

            if value == 'ﻣﺪﺭﺱ ﺍﻟﻤﺎﺩﺓ':
                cols[i] = 'teacherName'
            elif value == 'ﺍﻟﻤﺴﺠﻠﻴﻦ':
                cols[i] = 'registeredCount'
            elif value == 'ﺍﻻﻭﻗﺎﺕ':
                cols[i] = 'timeSlot'
            elif value == 'ﺍﻻﻳﺎﻡ':
                cols[i] = 'days'
            elif value == 'ﺍﻟﻤﺎﺩﺓ':
                cols[i] = 'courseName'
            elif value == 'ﺍﻟﻮﺣﺪﺍﺕ':
                cols[i] = 'credit'
            elif value == 'ﺍﻟﺸﻌﺒﺔ':
                cols[i] = 'group'
        df.columns = cols
        for i in range(1, len(df)):
            if df.loc[i, 'courseName'] == '':
                for j in str(df.loc[i, 'credit']).split('\n'):
                    if '-' in j:
                        df.loc[i, 'courseName'] = j

                    elif  2 <= len(j) <= 3:
                        df.loc[i, 'group'] = j
        df = df[['teacherName', 'registeredCount', 'timeSlot', 'days', 'courseName', 'group']]
        df = df[1:]
        df[['registeredCount','courseName']] = df[['registeredCount','courseName']].map(arabic_to_western)
        df[['timeSlot']] = df[['timeSlot']].map(extract_start_time)
        df[['days']] = df[['days']].map(get_earliest_day)
        df['group'] = df['group'].replace('', pd.NA).ffill()
        df['courseName'] = df['courseName'].replace('', pd.NA).ffill()
        df['dayPriority'] = df['days'].map(get_day_priority)


        dfs.append(df)


    merged_df = pd.concat(dfs, ignore_index=True)

    df = merged_df

    df.sort_values(by=['group', 'dayPriority', 'timeSlot'], inplace=True)
    df = clean_count(df)
    df['registeredCount'] = pd.to_numeric(df['registeredCount'], errors='coerce')
    df = df[df['registeredCount'].notna() & (df['registeredCount'] != 0)]
    df = df[df['teacherName'] != '']
    df = df.groupby(['courseName', 'group'], as_index=False).first()

    df['examid'] = df.apply(get_exam_id, axis=1)
    df = df.drop(columns=['days', 'group', 'dayPriority', 'timeSlot'])
    # Step 2: Group by 'courseName' and sum 'registeredCount'
    df = df.groupby('courseName', as_index=False).agg({
        'registeredCount': 'sum',
        'teacherName': lambda x: list(set(x)),  # collect all unique teacher names
        'examid': 'first'
    })
    df['courseName'] = df['courseName'].map(normalize_course_name)
    df = df[~df['courseName'].isin(['COCS-499', 'COCS-498', 'COIT-499', 'COIT-498', 'COIS-499', 'COIS-498', 'COCS-308'])]

    examid_map = {course: examid for examid, course in courses}

    df['examid'] = df.apply(
        lambda row: examid_map[row['courseName']] if row['courseName'] in examid_map else row['examid'],
        axis=1
    )
    df_main = df

    tables = camelot.read_pdf(file1, pages='3',flavor='lattice',process_background=True,line_scale=20)

    for i, table in enumerate(tables):
        df = table.df

        for j in range(len(df.iloc[0])):
            lstt = []
            for k in df[j][0].split('م'):
                clean_header(k.split('\n'))
                lstt.append(clean_header(k.split('\n')))
            df.loc[0,j] = lstt
        df.drop(df.columns[len(df.iloc[0])-1], axis=1, inplace=True)
        examid,datees,times= [],[],[]
        for row in range(1,len(df[0])):
            for col in range(len(df.iloc[0])):
                examid.append(df.loc[row,col])
                datees.append(df.loc[0,col])
                times.append(row)

        df_schedule = pd.DataFrame({
        'examid': examid,
        'examDate': datees,
        'timeSlot': times})
        df_main['examid'] = df_main['examid'].astype(int)
        df_schedule['examid'] = df_schedule['examid'].astype(int)
        df_merged = df_main.merge(df_schedule, on='examid', how='left')

        return create_exam_schedule_docx(df_merged)



